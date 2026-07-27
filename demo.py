from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from PyQt5.QtWidgets import (
    QApplication, QTableWidget, QTableWidgetItem, QWidget, QPushButton, QVBoxLayout, QHBoxLayout, QLabel,
    QInputDialog, QDialog, QComboBox, QGridLayout, QLineEdit, QMessageBox, QFileDialog, QTextEdit
)
from PyQt5.QtGui import QFont
import sys
import os
import numpy as np
import json

class BarChartCanvas(FigureCanvas):
    def __init__(self, parent=None, num_questions=21, answers=None):
        fig, self.ax = plt.subplots(figsize=(6, 6))
        super().__init__(fig)
        self.setParent(parent)
        self.num_questions = num_questions
        self.answers = answers or ["Low"] * self.num_questions
        #self.answers.append("Low")
        self.draw_chart()

    def draw_chart(self):
        self.ax.clear()
        y_labels = [f"Q{i+1}" for i in range(len(self.answers)-1)]
        y_labels.append("Prediction")
        y_pos = np.arange(len(self.answers))

        colors = {
            "G": "green",
            "Y": "yellow",
            "R": "red",
            "B": "blue",
            "P": "pink",
            "": "gray"   # unanswered stays gray
        }

        # Define colors for risk levels
        #colors = {'Low': 'yellow', 'Medium': 'orange', 'High': 'red'}
        bar_height = 1.0  # Full height per unit to eliminate gaps

        for i, risk in enumerate(self.answers):
            self.ax.barh(
                y_pos[i],
                10,  # Arbitrary same width
                height=bar_height,
                color=colors.get(risk, 'gray'),
                edgecolor='black',
                align='center'
            )

        self.ax.set_yticks(y_pos)
        self.ax.set_yticklabels(y_labels)
        self.ax.set_xlim(0, 10)
        self.ax.set_ylim(-0.5, len(self.answers) - 0.5)
        self.ax.invert_yaxis()  # Q1 at top
        #self.ax.set_title('Risk Levels')
        self.ax.set_xticks([])

        self.draw()

    """
    def draw_chart(self):
        self.ax.clear()
        color_map = {"Low": "yellow", "Medium": "orange", "High": "red"}
        for i, level in enumerate(self.answers):
            self.ax.barh(i, 1, color=color_map[level], edgecolor='black')
        self.ax.set_xlim(0, 1)
        self.ax.set_yticks(range(self.num_questions))
        self.ax.set_yticklabels([f"Q{i+1}" for i in range(self.num_questions)])
        self.ax.set_xticks([])
        self.ax.invert_yaxis()
        #self.ax.set_title("Risk Levels")
        self.draw()
    """

    def update_answers(self, new_answers):
        self.answers = new_answers
        self.draw_chart()

    def save_chart_image(self, filename="chart.png"):
        self.figure.savefig(filename)


class SurveyApp(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("VitaScreen")
        self.resize(1000,1000)

        script_dir = os.path.dirname(os.path.abspath(__file__))
        json_path = os.path.join(script_dir, 'VitaScreenDiabetesQuestions.json')
        with open(json_path, 'r', encoding='utf-8') as f:
            self.questions = json.load(f)["questions"]

        script_dir2 = os.path.dirname(os.path.abspath(__file__))
        json_path2 = os.path.join(script_dir2, 'sdataset.json')
        with open(json_path2, 'r', encoding='utf-8') as f:
            self.questions2 = json.load(f)["Questions"]

        script_dir3 = os.path.dirname(os.path.abspath(__file__))
        json_path3 = os.path.join(script_dir3, 'example.json')
        with open(json_path3, 'r', encoding='utf-8') as f:
            self.questions3 = json.load(f)["Questions"]

        self.answers = [""] * (len(self.questions) + 1)

        self.name_input = QLineEdit()
        self.health_id_input = QLineEdit()

        self.chart = BarChartCanvas(self, answers=self.answers)

        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        self.create_top_inputs()
        self.create_question_buttons()
        self.layout.addWidget(self.chart)
        self.create_action_buttons()



    def create_top_inputs(self):
        form_layout = QGridLayout()
        form_layout.addWidget(QLabel("Name:"), 0, 0)
        form_layout.addWidget(self.name_input, 0, 1)
        form_layout.addWidget(QLabel("Health Card ID:"), 1, 0)
        form_layout.addWidget(self.health_id_input, 1, 1)
        self.layout.addLayout(form_layout)

    def create_question_buttons(self):
        question_layout = QGridLayout()
        for i in range(len(self.questions)):
            button = QPushButton(f"Q{i+1}")
            button.clicked.connect(lambda _, idx=i: self.ask_question(idx))
            question_layout.addWidget(button, i // 5, i % 5)
        self.layout.addLayout(question_layout)

    def create_action_buttons(self):
        button_layout = QHBoxLayout()
        button_layout.addStretch()
        
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reset_answers)
        
        save_btn = QPushButton("Save")
        save_btn.clicked.connect(self.save_data)

        quit_btn = QPushButton("Quit")
        quit_btn.clicked.connect(self.quit_application)

        load_btn = QPushButton("Load")
        load_btn.clicked.connect(self.load_data)

        help_btn = QPushButton("Help")
        help_btn.clicked.connect(self.show_help)

        button_layout.addWidget(cancel_btn)
        button_layout.addWidget(save_btn)
        button_layout.addWidget(quit_btn)
        button_layout.addWidget(load_btn)
        button_layout.addWidget(help_btn)
        self.layout.addLayout(button_layout)

    def ask_question(self, idx):
        q_data = self.questions[idx]
        question_text = q_data.get("question", f"Question {idx+1}")

        options_dict = q_data.get("possible_answers", {})
        if not isinstance(options_dict, dict):
            print(f"Error: possible_answers is {type(options_dict)} for question {idx+1}")
            return

        # Build a list of answer strings for the dropdown
        options = []
        for val in options_dict.values():
            if isinstance(val, dict):
                options.append(val.get("answer", ""))
            elif isinstance(val, str):
                options.append(val)
        
        # Show dropdown
        choice, ok = QInputDialog.getItem(
            self,
            f"Answer for Q{idx+1}",
            question_text,
            options,
            editable=False
        )

        if ok and choice:
            # Map the selected answer back to its color code
            color = "G"  # default fallback
            for val in options_dict.values():
                if isinstance(val, dict) and val.get("answer") == choice:
                    color = val.get("color", "R")
                    break
                elif isinstance(val, str) and val == choice:
                    color = "G"

            # Save the color in answers
            self.answers[idx] = color

            # Update the chart
            self.chart.update_answers(self.answers)

    def reset_answers(self):
        self.answers = [""] * (len(self.questions)+1)
        self.chart.update_answers(self.answers)

    def save_report(self):
        name = self.name_input.text().strip()
        hcid = self.health_id_input.text().strip()
        if not name or not hcid:
            QMessageBox.warning(self, "Missing Info", "Please enter Name and Health Card ID before saving.")
            return

        save_path = QFileDialog.getSaveFileName(self, "Save Word Report", "", "Word Document (*.docx)")[0]
        if not save_path:
            return

        if not save_path.endswith(".docx"):
            save_path += ".docx"

        chart_img = "temp_chart.png"
        self.chart.save_chart_image(chart_img)

        doc = Document()
        doc.add_heading("Diabetic Risk Screening Report", 0)
        doc.add_paragraph(f"Name: {name}")
        doc.add_paragraph(f"Health Card ID: {hcid}")
        doc.add_picture(chart_img, width=Inches(5.5))
        doc.save(save_path)

        os.remove(chart_img)
        QMessageBox.information(self, "Saved", f"Report saved to {save_path}")

    def save_data(self):
        save_path = QFileDialog.getSaveFileName(self, "Save Data", "", "JSON Files (*.json)")[0]

        if not save_path:
            return

        if not save_path.endswith(".json"):
            save_path += ".json"

        data = {
            "name": self.name_input.text().strip(),
            "health_card_id": self.health_id_input.text().strip(),
            "answers": self.answers
        }

        with open(save_path, "w") as f:
            json.dump(data, f, indent=4)

        QMessageBox.information(self, "Saved", f"Report saved to {save_path}")

    def load_data(self):
        load_path = QFileDialog.getOpenFileName(self, "Load Data", "", "JSON Files (*.json)")[0]

        if not load_path:
            return

        with open(load_path, "r") as f:
            data = json.load(f)

        self.name_input.setText(data.get("name", ""))
        self.health_id_input.setText(data.get("health_card_id", ""))
        self.answers = data.get("answers", [""] * (len(self.questions)+1))
        self.chart.update_answers(self.answers)

    def quit_application(self):
        reply = QMessageBox.question(self, "Quit", "Do you want to save before quitting?", QMessageBox.Yes | QMessageBox.No | QMessageBox.Cancel)
        if reply == QMessageBox.Yes:
            self.save_report()
            QApplication.quit()
        elif reply == QMessageBox.No:
            QApplication.quit()
        # Cancel does nothing

    def load_report(self):
        load_path = QFileDialog.getOpenFileName(self, "Load Word Report", "", "Word Document (*.docx)")[0] #loads the report from a Word document
        if not load_path:
            return

        try:
            doc = Document(load_path)
            # Extract name and health card ID from the document
            name = "" # name and hcid are extracted from the Word document and set in the input fields -- currently blank if not found
            hcid = ""
            for para in doc.paragraphs:
                if para.text.startswith("Name:"):
                    name = para.text.replace("Name:", "").strip() # extracting the name from the Word document and setting it in the input field
                elif para.text.startswith("Health Card ID:"):
                    hcid = para.text.replace("Health Card ID:", "").strip() # extracting the health card ID from the Word document and setting it in the input field

            self.name_input.setText(name)
            self.health_id_input.setText(hcid)

            # Load chart image and update the chart
            chart_img = "temp_chart.png"
            for shape in doc.inline_shapes:
                if shape.type == 3:  # InlineShapeType.PICTURE
                    shape._inline.graphic.graphicData.pic.blipFill.blip.save(chart_img)
                    break

            if os.path.exists(chart_img):
                self.chart.figure.clear()
                img = plt.imread(chart_img)
                self.chart.ax.imshow(img)
                self.chart.ax.axis('off')
                self.chart.draw()
                os.remove(chart_img)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to load report: {e}")

    def show_help(self):
        dialog = QDialog(self)
        dialog.setWindowTitle("Survey Questions")
        dialog.resize(900, 900)

        layout = QVBoxLayout(dialog)

        self.tableWidget = QTableWidget()
        self.tableWidget.setRowCount(21)
        self.tableWidget.setColumnCount(4)
        self.tableWidget.setHorizontalHeaderLabels(["Question ID", "Question", "Feature Name", "Possible Values & Explanation"])

        for i,q in enumerate(self.questions):
            question_id = f"Q{i+1}"
            self.tableWidget.setItem(i, 0, QTableWidgetItem(question_id))
            self.tableWidget.setItem(i, 1, QTableWidgetItem(q.get('question', '')))
            self.tableWidget.setItem(i, 2, QTableWidgetItem(self.questions2.get(question_id, "")))
            self.tableWidget.setItem(i, 3, QTableWidgetItem(self.questions3.get(question_id, '')))

        layout.addWidget(self.tableWidget)
        
        dialog.exec_()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = SurveyApp()
    window.show()
    sys.exit(app.exec_())
